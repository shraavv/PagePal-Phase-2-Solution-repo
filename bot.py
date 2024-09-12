from docx import Document
import telebot
from telebot import types
from data import get_genre,preview_links
import os
from keys import BOT_TOKEN

bot = telebot.TeleBot(BOT_TOKEN)

user_input = {}
book_info = []
status = ''

file_path = 'reading_list.docx'

@bot.message_handler(commands=['start'])
def start_command(message):
    bot.reply_to(message, "Hello there! I am PagePal, your book companion.\nYou can ask me a specific genre and I will return a list of books! I can also provide you with preview links if available :)\nYou can also make a reading list of your own!")

@bot.message_handler(commands=['help'])
def help_command(message):
    bot.reply_to(message, "The commands are  \n1. start - Returns a welcome message \n2. book - Asks the user to enter the genre and it returns a CSV file with a list of book recommendations \n3. preview - Asks the user the book name and it redirects us to the preview link if available \n4. list - Asks the user for the book name which they want to add or delete from their reading list \n5. reading_list - Allows user to add a book, delete a book or view their reading list")

@bot.message_handler(commands=['book'])
def book_command(message):
    user_input[message.chat.id] = {'state': 'book'} 

    bot.reply_to(message, "Type in the genre name")

@bot.message_handler(commands=['preview'])
def preview_command(message):
    user_input[message.chat.id] = {'state': 'preview'}
    bot.reply_to(message, "Type in the book name for which you need the preview link")

@bot.message_handler(commands=['list'])
def bookname_command(message):
    user_input[message.chat.id] = {'state': 'list'}
    bot.reply_to(message, "Type in the book name")
    bot.reply_to(message, "After typing in the book name,execute /reading_list to perfom actions on your reading list")

def add_to_reading_list(book, link):
    global status
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs]
    if book not in paragraphs:
        doc.add_paragraph(book)
        doc.add_paragraph(link)
        doc.save(file_path)
        status = 'success'
    else:
        status='exists'
    book_info.clear()

def delete_from_reading_list(book):
    global status
    doc = Document(file_path)
    contents = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
    if book in contents:
        new_content = []
        i = 0
        while i < len(contents):
            if contents[i] != book:
                new_content.append(contents[i])
                new_content.append(contents[i + 1])
            i += 2

        doc = Document()
        for item in new_content:
            doc.add_paragraph(item)

        doc.save(file_path)
        
        status = 'success'
    else:
        status = 'nope'
    contents.clear()
    book_info.clear()
 
@bot.message_handler(commands=['reading_list'])
def question_commad(message):
    markup = types.InlineKeyboardMarkup(row_width=2)

    add_button = types.InlineKeyboardButton('Add to reading list', callback_data='add_button')
    read_button = types.InlineKeyboardButton('View your reading list', callback_data='read_button')
    delete_button = types.InlineKeyboardButton('Delete from your reading list', callback_data='delete_button')   

    markup.add(add_button, read_button, delete_button)

    bot.send_message(message.chat.id, 'You can perform the following functions on your reading list', reply_markup=markup)

@bot.callback_query_handler(func=lambda call:True)
def answer(callback):
    if callback.message:
        global status
        if callback.data == 'add_button':
            for i in book_info:
                    book = i[0]
                    link = i[1]
                    add_to_reading_list(book, link)
                    if status == 'success':
                        bot.send_message(callback.message.chat.id, 'Successfully added!')
                        status = ''
                    elif status == 'exists':
                        bot.send_message(callback.message.chat.id, 'Already exists!')
                        status = ''
        elif callback.data == 'read_button':
            if os.path.exists(file_path):
                bot.send_document(callback.message.chat.id, open(file_path, 'rb'))
            else:
                bot.send_message(callback.message.chat.id, "The reading list is empty :(")
        elif callback.data == 'delete_button':
            for i in book_info:
                    book = i[0]
                    delete_from_reading_list(book)
                    if status == 'success':
                        bot.send_message(callback.message.chat.id, 'Successfully deleted!')
                        status = ''
                    elif status == 'nope' :
                        bot.send_message(callback.message.chat.id, 'Does not exist!')
                        status = ''
                        
@bot.message_handler(func=lambda msg: True)
def echo_input(message):
    chat_id = message.chat.id
    if chat_id in user_input:
        state = user_input[chat_id]['state']
        if state == 'book':
            genre = message.text.lower()
            get_genre(genre)
            with open('data.csv', 'rb') as file:
                file_size = os.path.getsize('data.csv')
                if file_size > 0:
                    bot.reply_to(message,"Here's your CSV file with a list of book recommendations")
                    bot.send_document(chat_id, file)
                else:
                    bot.reply_to(message,"No recommendations :( ")
            del user_input[chat_id]

        elif state == 'preview':
            book_name = message.text.lower()
            link = preview_links(book_name)
            if link.startswith("http"):
                bot.reply_to(message,"Here's the preview link")
                bot.reply_to(message,link)
            else:
                bot.reply_to(message,"Preview not available :/")
            del user_input[chat_id]
        
        elif state == 'list':
            book_name = message.text.lower()
            link = preview_links(book_name)
            if link.startswith("http"):
                book_info.append([book_name, link])
            else:
                bot.reply_to(message,"Book not available :/")
            del user_input[chat_id]
    
    else:
        bot.reply_to(message, "Sorry could not process :(")

bot.infinity_polling()